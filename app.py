import streamlit as st
import anthropic
import json
from playwright.sync_api import sync_playwright
import tempfile
import os
import fitz  # PyMuPDF
from docx import Document as DocxDoc
import subprocess
import os
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
try:
    subprocess.run(
        ["python", "-m", "playwright", "install", "chromium"],
        check=False
    )
except Exception:
    pass
# ── Page config ───────────────────────────────────────────────
st.set_page_config(
    page_title="Subject Overview Generator · UNext",
    page_icon="📚",
    layout="centered",
)

st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500&display=swap');
  html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
  .block-container { max-width: 680px; padding-top: 2.5rem; }
  .stButton > button {
    background: #0D2D6B; color: white; border: none;
    border-radius: 8px; padding: 12px 32px;
    font-size: 15px; font-weight: 500; width: 100%;
    transition: background 0.2s;
  }
  .stButton > button:hover { background: #1A4099; border: none; }
  .stDownloadButton > button {
    background: #0F6E56; color: white; border: none;
    border-radius: 8px; padding: 12px 32px;
    font-size: 15px; font-weight: 500; width: 100%;
  }
  .stDownloadButton > button:hover { background: #0A5540; border: none; }
  .eyebrow {
    font-size: 11px; font-weight: 500; letter-spacing: .16em;
    text-transform: uppercase; color: #F47920; margin-bottom: 4px;
  }
  .divider { height: 2px; background: linear-gradient(90deg,#0D2D6B 52%,#0F6E56 76%,#F47920 100%);
    border-radius: 2px; margin: 24px 0; }
</style>
""", unsafe_allow_html=True)

# ── Inline SVG icon paths ─────────────────────────────────────
ICONS = {
  "file-text":    '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M14 3v4a1 1 0 0 0 1 1h4"/><path d="M17 21h-10a2 2 0 0 1 -2 -2v-14a2 2 0 0 1 2 -2h7l5 5v11a2 2 0 0 1 -2 2"/><path d="M9 9l1 0"/><path d="M9 13l6 0"/><path d="M9 17l6 0"/>',
  "bulb":         '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 12h1m8 -9v1m8 8h1m-15.4 -6.4l.7 .7m12.1 -.7l-.7 .7"/><path d="M9 16a5 5 0 1 1 6 0a3.5 3.5 0 0 0 -1 3a2 2 0 0 1 -4 0a3.5 3.5 0 0 0 -1 -3"/><path d="M9.7 17l4.6 0"/>',
  "gauge":        '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 12a9 9 0 1 0 18 0a9 9 0 1 0 -18 0"/><path d="M11 12a1 1 0 1 0 2 0a1 1 0 1 0 -2 0"/><path d="M13.41 10.59l2.59 -2.59"/><path d="M7 12a5 5 0 0 1 5 -5"/>',
  "rocket":       '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M4 13a8 8 0 0 1 7 7a6 6 0 0 0 3 -5a9 9 0 0 0 6 -8a3 3 0 0 0 -3 -3a9 9 0 0 0 -8 6a6 6 0 0 0 -5 3"/><path d="M7 14a6 6 0 0 0 -3 6a6 6 0 0 0 6 -3"/><path d="M14 9a1 1 0 1 0 2 0a1 1 0 1 0 -2 0"/>',
  "target":       '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M11 12a1 1 0 1 0 2 0a1 1 0 1 0 -2 0"/><path d="M7 12a5 5 0 1 0 10 0a5 5 0 1 0 -10 0"/><path d="M3 12a9 9 0 1 0 18 0a9 9 0 1 0 -18 0"/>',
  "cpu":          '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M5 6a1 1 0 0 1 1 -1h12a1 1 0 0 1 1 1v12a1 1 0 0 1 -1 1h-12a1 1 0 0 1 -1 -1l0 -12"/><path d="M9 9h6v6h-6l0 -6"/><path d="M3 10h2"/><path d="M3 14h2"/><path d="M10 3v2"/><path d="M14 3v2"/><path d="M21 10h-2"/><path d="M21 14h-2"/><path d="M14 21v-2"/><path d="M10 21v-2"/>',
  "brain":        '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M15.5 13a3.5 3.5 0 0 0 -3.5 3.5v1a3.5 3.5 0 0 0 7 0v-1.8"/><path d="M8.5 13a3.5 3.5 0 0 1 3.5 3.5v1a3.5 3.5 0 0 1 -7 0v-1.8"/><path d="M17.5 16a3.5 3.5 0 0 0 0 -7h-.5"/><path d="M19 9.3v-2.8a3.5 3.5 0 0 0 -7 0"/><path d="M6.5 16a3.5 3.5 0 0 1 0 -7h.5"/><path d="M5 9.3v-2.8a3.5 3.5 0 0 1 7 0v10"/>',
  "chart-line":   '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M4 19l16 0"/><path d="M4 15l4 -6l4 2l4 -5l4 4"/>',
  "chart-bar":    '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 13a1 1 0 0 1 1 -1h4a1 1 0 0 1 1 1v6a1 1 0 0 1 -1 1h-4a1 1 0 0 1 -1 -1z"/><path d="M15 9a1 1 0 0 1 1 -1h4a1 1 0 0 1 1 1v10a1 1 0 0 1 -1 1h-4a1 1 0 0 1 -1 -1z"/><path d="M9 5a1 1 0 0 1 1 -1h4a1 1 0 0 1 1 1v14a1 1 0 0 1 -1 1h-4a1 1 0 0 1 -1 -1z"/><path d="M4 20h14"/>',
  "chart-scatter":'<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 3v18h18"/><path d="M8 15.015v.015"/><path d="M16 16.015v.015"/><path d="M8 7.03v.015"/><path d="M12 11.03v.015"/><path d="M19 11.03v.015"/>',
  "database":     '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M4 6a8 3 0 1 0 16 0a8 3 0 1 0 -16 0"/><path d="M4 6v6a8 3 0 0 0 16 0v-6"/><path d="M4 12v6a8 3 0 0 0 16 0v-6"/>',
  "trending-up":  '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 17l6 -6l4 4l8 -8"/><path d="M14 7l7 0l0 7"/>',
  "letter-case":  '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M14 15.5a3.5 3.5 0 1 0 7 0a3.5 3.5 0 1 0 -7 0"/><path d="M3 19v-10.5a3.5 3.5 0 0 1 7 0v10.5"/><path d="M3 13h7"/><path d="M21 12v7"/>',
  "message-dots": '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M12 11v.01"/><path d="M8 11v.01"/><path d="M16 11v.01"/><path d="M18 4a3 3 0 0 1 3 3v8a3 3 0 0 1 -3 3h-5l-5 3v-3h-2a3 3 0 0 1 -3 -3v-8a3 3 0 0 1 3 -3l12 0"/>',
  "robot":        '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M6 6a2 2 0 0 1 2 -2h8a2 2 0 0 1 2 2v4a2 2 0 0 1 -2 2h-8a2 2 0 0 1 -2 -2z"/><path d="M12 2v2"/><path d="M9 12v9"/><path d="M15 12v9"/><path d="M5 16l4 -2"/><path d="M15 14l4 2"/><path d="M9 18h6"/><path d="M10 8v.01"/><path d="M14 8v.01"/>',
  "report-analytics": '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M9 5h-2a2 2 0 0 0 -2 2v12a2 2 0 0 0 2 2h10a2 2 0 0 0 2 -2v-12a2 2 0 0 0 -2 -2h-2"/><path d="M9 5a2 2 0 0 1 2 -2h2a2 2 0 0 1 2 2a2 2 0 0 1 -2 2h-2a2 2 0 0 1 -2 -2"/><path d="M9 17v-5"/><path d="M12 17v-1"/><path d="M15 17v-3"/>',
  "flask":        '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M9 3l6 0"/><path d="M10 9l4 0"/><path d="M10 3v6l-4 11a.7 .7 0 0 0 .5 1h11a.7 .7 0 0 0 .5 -1l-4 -11v-6"/>',
  "briefcase":    '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 9a2 2 0 0 1 2 -2h14a2 2 0 0 1 2 2v9a2 2 0 0 1 -2 2h-14a2 2 0 0 1 -2 -2z"/><path d="M8 7v-2a2 2 0 0 1 2 -2h4a2 2 0 0 1 2 2v2"/><path d="M12 12l0 .01"/><path d="M3 13a20 20 0 0 0 18 0"/>',
  "microscope":   '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M5 21h14"/><path d="M6 18h2"/><path d="M7 18v3"/><path d="M9 11l3 3l6 -6l-3 -3z"/><path d="M10.5 12.5l-1.5 1.5"/><path d="M17 3l3 3"/><path d="M12 21a6 6 0 0 0 3.715 -10.712"/>',
  "tool":         '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M7 10h3v-3l-3.5 -3.5a6 6 0 0 1 8 8l6 6a2 2 0 0 1 -3 3l-6 -6a6 6 0 0 1 -8 -8z"/>',
  "shield-check": '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M11.46 20.846a12 12 0 0 1 -7.96 -14.846a12 12 0 0 0 8.5 -3a12 12 0 0 0 8.5 3a12 12 0 0 1 -.09 7.06"/><path d="M15 19l2 2l4 -4"/>',
  "code":         '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M7 8l-4 4l4 4"/><path d="M17 8l4 4l-4 4"/><path d="M14 4l-4 16"/>',
  "adjustments-horizontal": '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M12 6a2 2 0 1 0 4 0a2 2 0 1 0 -4 0"/><path d="M4 6l8 0"/><path d="M16 6l4 0"/><path d="M6 12a2 2 0 1 0 4 0a2 2 0 1 0 -4 0"/><path d="M4 12l2 0"/><path d="M10 12l10 0"/><path d="M15 18a2 2 0 1 0 4 0a2 2 0 1 0 -4 0"/><path d="M4 18l11 0"/><path d="M19 18l1 0"/>',
  "hierarchy-2":  '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M10 3h4v4h-4z"/><path d="M3 17h4v4h-4z"/><path d="M17 17h4v4h-4z"/><path d="M7 17l5 -4l5 4"/><path d="M12 7l0 6"/>',
  "scale":        '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M7 20l10 0"/><path d="M6 6l6 -1l6 1"/><path d="M12 3l0 17"/><path d="M9 12l-3 -6l-3 6a3 3 0 0 0 6 0"/><path d="M21 12l-3 -6l-3 6a3 3 0 0 0 6 0"/>',
  "book":         '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 19a9 9 0 0 1 9 0a9 9 0 0 1 9 0"/><path d="M3 6a9 9 0 0 1 9 0a9 9 0 0 1 9 0"/><path d="M3 6l0 13"/><path d="M12 6l0 13"/><path d="M21 6l0 13"/>',
  "search":       '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 10a7 7 0 1 0 14 0a7 7 0 1 0 -14 0"/><path d="M21 21l-6 -6"/>',
  "star":         '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M12 17.75l-6.172 3.245l1.179 -6.873l-5 -4.867l6.9 -1l3.086 -6.253l3.086 6.253l6.9 1l-5 4.867l1.179 6.873z"/>',
  "map":          '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M3 7l6 -3l6 3l6 -3v13l-6 3l-6 -3l-6 3v-13"/><path d="M9 4v13"/><path d="M15 7v13"/>',
  "checks":       '<path stroke="none" d="M0 0h24v24H0z" fill="none"/><path d="M7 12l5 5l10 -10"/><path d="M2 12l5 5m5 -5l5 -5"/>',
}

VALID_ICONS = list(ICONS.keys())
def assign_icons(data):

    ICON_MAP = {

        # International Relations / Political Science
        "foreign": "map",
        "international": "map",
        "global": "map",
        "diplomacy": "message-dots",
        "communication": "message-dots",
        "policy": "report-analytics",
        "governance": "scale",
        "government": "scale",
        "law": "scale",
        "ethics": "scale",

        # Strategy & Management
        "strategy": "target",
        "strategic": "target",
        "leadership": "briefcase",
        "management": "hierarchy-2",
        "project": "hierarchy-2",

        # Analytics & Research
        "analysis": "chart-line",
        "analyse": "chart-line",
        "research": "search",
        "data": "database",
        "statistics": "chart-scatter",
        "probability": "chart-scatter",

        # Finance & Economics
        "finance": "trending-up",
        "financial": "trending-up",
        "investment": "trending-up",
        "economics": "chart-bar",
        "economic": "chart-bar",
        "accounting": "checks",
        "audit": "checks",
        "auditing": "checks",

        # Technology
        "technology": "cpu",
        "digital": "cpu",
        "artificial intelligence": "robot",
        "machine learning": "robot",
        "ai": "robot",
        "coding": "code",
        "programming": "code",

        # Literature & Language
        "literature": "book",
        "literary": "book",
        "novel": "book",
        "fiction": "book",
        "writing": "file-text",
        "journalism": "file-text",
        "media": "file-text",
        "language": "letter-case",
        "linguistics": "letter-case",

        # Science
        "science": "microscope",
        "scientific": "microscope",

        # Security
        "security": "shield-check",
        "risk": "shield-check",

        # Innovation
        "innovation": "bulb"
    }

    for section in ["outcomes", "competencies", "roles"]:

        for item in data.get(section, []):

            if section == "roles":
                text = (
                    item.get("title", "") + " " +
                    item.get("body", "")
                ).lower()
            else:
                text = (
                    item.get("keyword", "") + " " +
                    item.get("rest", "") + " " +
                    item.get("body", "")
                ).lower()

            for keyword, icon in ICON_MAP.items():
                if keyword in text:
                    item["icon"] = icon
                    break

    return data
# ── Claude prompt ─────────────────────────────────────────────
SYSTEM_PROMPT = f"""You are a curriculum design assistant. Analyse the provided syllabus and return ONLY valid JSON — no markdown fences, no explanation, just the raw JSON object.

JSON structure:
{{
  "subjectLine1": "first part of subject name",
  "subjectLine2": "second part in italic orange (or empty string if single line)",
  "outcomes": [
    {{ "icon": "icon-name", "keyword": "BloomsVerb", "rest": " rest of title text", "body": "one clear sentence" }}
  ],
  "competencies": [
    {{ "icon": "icon-name", "keyword": "Key Word", "rest": " rest of title text", "body": "one clear sentence" }}
  ],
  "roles": [
    {{ "icon": "icon-name", "title": "Job Title", "body": "specific real-world task using subject skills" }}
  ]
}}

Rules:
- outcomes: exactly 4 items.
- Each keyword must be a Bloom's Taxonomy verb: Analyse, Build, Design, Evaluate, Apply, Create, Interpret, or Construct.
- Outcome titles should be concise and action-oriented.
- Outcome body text should be 10-15 words maximum.
- Focus on practical learner capability rather than academic description.

- competencies: exactly 4 items.
- Competency titles should be short, professional skill statements.
- Competency body text should be 10-15 words maximum.
- Avoid generic educational language.

- roles: exactly 4 items.
- Use realistic industry job titles.
- Job role descriptions must describe what the professional actually does day-to-day.
- Body text should be 15-20 words maximum.
- Avoid textbook explanations.

Writing Style:
- Be concise and professional.
- Avoid phrases such as "students will learn", "students examine", "students understand".
- Write in a modern, industry-oriented style.
- Use strong action verbs.
- Prefer practical outcomes over academic descriptions.
- Every sentence should be easily scannable on a poster.
- Minimise unnecessary words.
- Return ONLY the JSON object. No markdown. No extra text."""


# ── File reading ───────────────────────────────────────────────
def read_file(uploaded_file):
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "pdf":
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    elif ext == "docx":
        doc = DocxDoc(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")


# ── API call ───────────────────────────────────────────────────
def extract_content(text, api_key):
    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": f"Syllabus:\n\n{text[:12000]}"}],
    )
    raw = response.content[0].text.strip()
    clean = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(clean)


# ── HTML poster builder ────────────────────────────────────────
def svg_icon(name, color, size=22):
    paths = ICONS.get(name, ICONS["star"])
    return f'''<svg viewBox="0 0 24 24" width="{size}" height="{size}"
      fill="none" stroke="{color}" stroke-width="1.6"
      stroke-linecap="round" stroke-linejoin="round">{paths}</svg>'''


def build_poster_html(data):
    cols = [
        {"key": "outcomes",     "label": "Learning Outcomes",     "bg": "#EBF0FA", "accent": "#0D2D6B"},
        {"key": "competencies", "label": "Competencies / Skills",  "bg": "#E1F5EE", "accent": "#0F6E56"},
        {"key": "roles",        "label": "Job Roles",              "bg": "#FEF1E6", "accent": "#F47920"},
    ]

    def item_html(item, accent, col_bg, is_role, is_last):
        icon  = svg_icon(item.get("icon", "star"), accent, 22)
        if is_role:
            title_html = f'<strong style="color:{accent};font-weight:600">{item.get("title","")}</strong>'
        else:
            title_html = (f'<strong style="color:{accent};font-weight:600">{item.get("keyword","")}</strong>'
                          f'{item.get("rest","")}')
        border = "none" if is_last else f"0.5px solid rgba(0,0,0,0.09)"
        return f"""
        <div style="display:flex;gap:14px;align-items:flex-start;
                    padding:{'0 0 20px' if not is_last else '20px 0 0'};
                    border-bottom:{border};">
          <div style="flex-shrink:0;width:40px;height:40px;border-radius:10px;
                      background:{col_bg};display:flex;align-items:center;
                      justify-content:center;margin-top:1px;">{icon}</div>
          <div style="flex:1">
            <p style="font-size:15px;font-weight:500;color:#111827;
                      line-height:1.35;margin:0 0 5px">{title_html}</p>
            <p style="font-size:13px;font-weight:300;color:#4B5563;
                      line-height:1.72;margin:0">{item.get("body","")}</p>
          </div>
        </div>"""

    col_blocks = ""
    for ci, col in enumerate(cols):
        items = data.get(col["key"], [])
        items_html = ""
        for ii, item in enumerate(items):
            is_last = ii == len(items) - 1
            items_html += item_html(item, col["accent"], col["bg"],
                                    col["key"] == "roles", is_last)
        border_right = "0.5px solid rgba(0,0,0,0.09)" if ci < 2 else "none"
        pad_left  = "24px"
        pad_right = "24px"        
        col_blocks += f"""
        <div style="padding-left:{pad_left};padding-right:{pad_right};
                    border-right:{border_right};">
          <div style="display:flex;align-items:center;gap:10px;padding:10px 14px;
                      border-radius:7px;background:{col["bg"]};margin-bottom:28px;min-height:46px;">
            <div style="width:8px;height:8px;border-radius:50%;
                        background:{col["accent"]};flex-shrink:0;"></div>
            <span style="font-size:10.5px;font-weight:600;letter-spacing:.12em;
                         text-transform:uppercase;color:{col["accent"]};line-height:1.3;">
              {col["label"]}</span>
          </div>
          {items_html}
        </div>"""

    line2 = data.get("subjectLine2", "")
    title_html = (
        f'<span style="color:#0D2D6B">{data.get("subjectLine1","")}</span>'
        + (f'<br><em style="font-style:italic;color:#F47920">{line2}</em>' if line2 else "")
    )

    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:ital,wght@0,300;0,400;0,500;1,300&family=DM+Serif+Display:ital@0;1&display=swap" rel="stylesheet">
<style>*{{box-sizing:border-box;margin:0;padding:0}}</style>
</head><body>
<div style="font-family:'DM Sans',sans-serif;width:1000px;padding:62px 60px 70px;
            background:#fff;position:relative;overflow:hidden;">

  <!-- blobs -->
  <div style="position:absolute;border-radius:50%;opacity:.05;width:280px;height:280px;
              background:#F47920;top:-90px;right:-80px;pointer-events:none;"></div>
  <div style="position:absolute;border-radius:50%;opacity:.05;width:200px;height:200px;
              background:#0D2D6B;bottom:-60px;left:-60px;pointer-events:none;"></div>
  <div style="position:absolute;border-radius:50%;opacity:.05;width:130px;height:130px;
              background:#0F6E56;bottom:80px;right:40px;pointer-events:none;"></div>

  <p style="font-size:12px;font-weight:500;letter-spacing:.17em;text-transform:uppercase;
            color:#F47920;margin-bottom:12px;position:relative;">Subject Overview</p>

  <h1 style="font-family:'DM Serif Display',serif;font-size:46px;font-weight:400;
             line-height:1.08;margin-bottom:28px;position:relative;">{title_html}</h1>

  <div style="height:3px;border-radius:2px;margin-bottom:44px;
              background:linear-gradient(90deg,#0D2D6B 52%,#0F6E56 76%,#F47920 100%);"></div>

  <div style="display:grid;grid-template-columns:repeat(3,1fr);column-gap:8px;">
       {col_blocks}
  </div>
</div>
</body></html>"""


# ── Render PNG via playwrite ─────────────────────────────────────
def render_png(html):
    with tempfile.TemporaryDirectory() as tmpdir:
        html_file = os.path.join(tmpdir, "poster.html")

        with open(html_file, "w", encoding="utf-8") as f:
            f.write(html)

        with sync_playwright() as p:
            browser = p.chromium.launch(
                executable_path="/usr/bin/chromium",
                headless=True,
                args=[
                    "--no-sandbox",
                    "--disable-dev-shm-usage",
                    "--disable-gpu"
                ]
            )

            page = browser.new_page(
                viewport={"width": 1000, "height": 1400},
                device_scale_factor=3
            )

            page.goto(f"file://{html_file}")

            png_bytes = page.screenshot(
                full_page=True,
                type="png"
            )

            browser.close()

        return png_bytes
# ── Main UI ────────────────────────────────────────────────────
def build_docx(data):

    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run(data.get("subjectLine1", ""))
    run.bold = True
    run.font.size = Pt(24)

    if data.get("subjectLine2"):
        title.add_run("\n")
        run2 = title.add_run(data["subjectLine2"])
        run2.italic = True
        run2.font.size = Pt(24)
        run2.font.color.rgb = RGBColor(244, 121, 32)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.autofit = True

    sections = [
        ("Learning Outcomes", data["outcomes"]),
        ("Competencies / Skills", data["competencies"]),
        ("Job Roles", data["roles"])
    ]

    for i, (heading, items) in enumerate(sections):

        cell = table.rows[0].cells[i]

        p = cell.paragraphs[0]
        run = p.add_run(heading)
        run.bold = True

        p.add_run("\n\n")

        for item in items:

            if "title" in item:
                title_text = item["title"]
            else:
                title_text = (
                    item.get("keyword", "") +
                    item.get("rest", "")
                )

            r = p.add_run(title_text + "\n")
            r.bold = True

            p.add_run(item.get("body", "") + "\n\n")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
def main():
    st.markdown('<p class="eyebrow">UNext Learning</p>', unsafe_allow_html=True)
    st.title("Subject Overview Generator")
    st.markdown(
        "Upload any subject syllabus and get a beautifully designed, "
        "ready-to-download overview image in seconds."
    )
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # Load API key from Streamlit secrets
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        st.error(
            "⚠ API key not found. "
            "Please add `ANTHROPIC_API_KEY` to your Streamlit secrets "
            "(Settings → Secrets in the Streamlit Cloud dashboard)."
        )
        st.stop()

    uploaded_file = st.file_uploader(
        "Upload your syllabus",
        type=["pdf", "docx", "txt"],
        help="Supports PDF, Word (.docx), and plain text files",
    )

    if uploaded_file:
        st.caption(f"📄 {uploaded_file.name} uploaded")

        if st.button("Generate Subject Overview →"):
            with st.status("Working on it…", expanded=True) as status:
                st.write("📖 Reading your syllabus…")
                try:
                    text = read_file(uploaded_file)
                except Exception as e:
                    st.error(f"Could not read file: {e}")
                    st.stop()

                st.write("🧠 Extracting outcomes, competencies and job roles…")
                try:
                    data = extract_content(text, api_key)
                    data = assign_icons(data)                    
                except Exception as e:
                    st.error(f"Content extraction failed: {e}")
                    st.stop()

                st.write("🎨 Generating poster…")
                try:
                    html    = build_poster_html(data)
                    png_img = render_png(html)
                    docx_file = build_docx(data)
                except Exception as e:
                    st.error(f"Image render failed: {e}")
                    st.stop()

                status.update(label="✅ Done!", state="complete")

            st.image(png_img, use_container_width=True)

            subject_slug = (
                data.get("subjectLine1", "subject") + "_" + data.get("subjectLine2", "")
            ).strip("_").replace(" ", "_").replace("&", "and")

            st.download_button(
                label="⬇  Download PNG",
                data=png_img,
                file_name=f"subject_overview_{subject_slug}.png",
                mime="image/png",
            )
            st.download_button(
                label = "⬇ Download Editable Word Document",
                data = docx_file,
                file_name=f"subject_overview_{subject_slug}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

if __name__ == "__main__":
    main()
